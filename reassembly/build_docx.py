import os
import sqlite3
import re
from docx import Document
from .fonts import set_devanagari_font

def export_marathi_docx(db_path: str, output_path: str, font_name="Nirmala UI"):
    """
    Reads the final translated paragraphs from the database and constructs a Word document,
    applying the Devanagari font fix to ensure proper rendering.
    """
    print(f"Building final Marathi document at {output_path}...")
    doc = Document()
    
    conn = sqlite3.connect(db_path)
    c = conn.cursor()
    
    # Fetch all finalized translations in sequence order
    try:
        c.execute("SELECT element_type, final_translation, source_text FROM paragraphs ORDER BY seq_order ASC")
        rows = c.fetchall()
    except sqlite3.OperationalError:
        print("Database not initialized or empty.")
        return
    
    for row in rows:
        element_type, marathi_text, english_text = row
        
        # If it bypassed translation (e.g., code_block, table, equation) it might only have source_text
        text_to_write = marathi_text if marathi_text else english_text
        if not text_to_write:
            continue
            
        # Sanitize string to prevent python-docx XML ValueError
        text_to_write = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]', '', text_to_write)
            
        p = doc.add_paragraph()
        run = p.add_run(text_to_write)
        
        # Apply the XML font fix
        if element_type not in ["code_block", "inline_code"]:
            set_devanagari_font(run, font_name=font_name, size_pt=12)
        else:
            # Code blocks use monospace font
            run.font.name = "Consolas"
            
        # Simple formatting based on element type
        if element_type == "heading":
            run.bold = True
            
    doc.save(output_path)
    print(f"Document saved successfully: {output_path}")

if __name__ == "__main__":
    print("Reassembly module ready.")
