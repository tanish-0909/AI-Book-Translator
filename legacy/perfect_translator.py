import os
import re
import warnings
warnings.filterwarnings('ignore')

from pdf2docx import Converter
from docx import Document
import torch
from docx.oxml.ns import qn

from translation.draft import load_draft_model
from reassembly.fonts import set_devanagari_font

def sanitize(text):
    return re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]', '', text)

def run_perfect_pipeline():
    pdf_file = "input/Book.pdf"
    temp_docx = "workdir/layout_temp.docx"
    out_docx = "output/book_marathi.docx"
    
    os.makedirs("workdir", exist_ok=True)
    os.makedirs("output", exist_ok=True)
    
    print("1. Performing pixel-perfect layout preservation (PDF -> DOCX)...")
    if not os.path.exists(temp_docx):
        cv = Converter(pdf_file)
        cv.convert(temp_docx)
        cv.close()
    
    print("2. Extracting structural nodes...")
    doc = Document(temp_docx)
    nodes_to_translate = []
    
    # Collect nodes from paragraphs
    for p in doc.paragraphs:
        if p.text.strip():
            nodes_to_translate.append(p)
            
    # Collect nodes from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if p.text.strip():
                        nodes_to_translate.append(p)
                        
    print(f"Found {len(nodes_to_translate)} structural elements.")
    
    # Load translation engine
    print("3. Hot-swapping text with IndicTrans2...")
    tokenizer, model, ip = load_draft_model()
    
    batch_size = 8
    for i in range(0, len(nodes_to_translate), batch_size):
        batch_nodes = nodes_to_translate[i:i+batch_size]
        batch_texts = [n.text.strip() for n in batch_nodes]
        
        prepped = ip.preprocess_batch(batch_texts, src_lang="eng_Latn", tgt_lang="mar_Deva")
        inputs = tokenizer(prepped, padding=True, truncation=True, max_length=1024, return_tensors="pt").to(model.device)
        
        with torch.no_grad():
            generated = model.generate(**inputs, max_length=1024, num_beams=5)
            
        decoded = tokenizer.batch_decode(generated, skip_special_tokens=True)
        postprocessed = ip.postprocess_batch(decoded, lang="mar_Deva")
        
        for node, translated_text in zip(batch_nodes, postprocessed):
            # Preserve paragraph format but swap the text
            node.clear()
            run = node.add_run(sanitize(translated_text))
            
            # Complex Devanagari script fixing
            set_devanagari_font(run, font_name="Mangal")
            
    print(f"4. Saving perfectly formatted document to {out_docx}")
    doc.save(out_docx)
    print("DONE! Check the output document.")

if __name__ == "__main__":
    run_perfect_pipeline()
